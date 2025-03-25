import streamlit as st
from PyPDF2 import PdfReader
from langchain.docstore.document import Document
from langchain.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.summarize import load_summarize_chain
from docx import Document as DocxDocument
from docx.shared import Pt
import re
import os
import langdetect

st.title("Circulars Summary Generator")

api_key = st.text_input("Enter your OpenAI API Key", type="password")
st.caption("Your API key should start with 'sk-' and will not be stored")

uploaded_files = st.file_uploader("Upload PDF files", type="pdf", accept_multiple_files=True)

summarize_button = st.button("Summarize")

def extract_english_text(text):
    try:
        words = re.findall(r'\b\w+\b', text)
        
        english_words = []
        for word in words:
            try:
                if len(word) > 1:
                    lang = langdetect.detect(word)
                    if lang == 'en':
                        english_words.append(word)
            except langdetect.lang_detect_exception.LangDetectException:
                continue
        
        return ' '.join(english_words)
    
    except Exception as e:
        st.warning(f"Language error: {e}")
        return text

def chunk_document(text, chunk_size=8000, chunk_overlap=500):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    
    return text_splitter.split_text(text)

if summarize_button and uploaded_files and api_key:
    if not api_key.startswith("sk-"):
        st.error("Invalid API key format. OpenAI API keys should start with 'sk-'")
    else:
        try:
            llm = ChatOpenAI(
                openai_api_key=api_key,
                model_name="gpt-4o-2024-08-06",
                temperature=0.2,
                top_p=0.2
            )
            
            map_prompt = PromptTemplate(
                input_variables=["text"],
                template="""Read and summarize the following content chunk in your own words, highlighting the main ideas, purpose, and important insights without including direct phrases or sentences from the original text in 10 bullet points.\n\n{text}
                """
            )
            
            combine_prompt = PromptTemplate(
                input_variables=["text"],
                template="""
                Combine the following individual chunk summaries into a cohesive, insightful summary for the entire document. Ensure that it captures the core themes and purpose across all chunks in 15 bullet points:
                Each summary for a document should start with the document name (without extensions like .pdf or .docx).
                Each summary should have a heading named, "Key Pointers:"
                Combine the following individual summaries into a cohesive, insightful summary. Ensure that it is concise, capturing the core themes and purpose of the entire document in 15 bullet points:\n\n{text}
                """
            )
            
            # Create Word document for summaries
            doc = DocxDocument()
            
            # Add Consolidated Overview Summary header
            consolidated_overview_heading = doc.add_paragraph()
            consolidated_overview_run = consolidated_overview_heading.add_run("Consolidated Overview Summary")
            consolidated_overview_run.bold = True
            consolidated_overview_run.font.size = Pt(16)
            doc.add_paragraph()  # Add an empty line after header
            
            with st.spinner("Processing documents"):
                all_document_summaries = []
                partially_filtered_files = []
                
                for uploaded_file in uploaded_files:
                    file_progress = st.empty()
                    file_progress.text(f"Processing {uploaded_file.name}...")
                    
                    pdf = PdfReader(uploaded_file)
                    text = ''
                    for page in pdf.pages:
                        content = page.extract_text()
                        if content:
                            text += content + "\n"
                    
                    filtered_text = extract_english_text(text)
                    
                    if not filtered_text.strip():
                        st.warning(f"No English text found in {uploaded_file.name}")
                        continue
                    if len(filtered_text) < len(text) * 0.5:
                        partially_filtered_files.append(uploaded_file.name)
                    
                    text_chunks = chunk_document(filtered_text)
                    docs = [Document(page_content=chunk) for chunk in text_chunks]
                    map_reduce_chain = load_summarize_chain(
                        llm,
                        chain_type="map_reduce",
                        map_prompt=map_prompt,
                        combine_prompt=combine_prompt
                    )
                    
                    output_summary = map_reduce_chain.invoke(docs)
                    summary = output_summary['output_text']
                    
                    # Add to Word document with formatting similar to base code
                    document_name = uploaded_file.name.replace(".pdf", "")
                    
                    # Add bold document name
                    doc_heading = doc.add_paragraph()
                    heading_run = doc_heading.add_run(f"Document Name: {document_name}")
                    heading_run.bold = True
                    heading_run.font.size = Pt(14)
                    
                    key_pointers_heading = doc.add_paragraph()
                    key_pointers_run = key_pointers_heading.add_run("Key Pointers:")
                    key_pointers_run.bold = True
                    key_pointers_run.font.size = Pt(12)
                    
                    # Parse and extract bullet points
                    bullet_pattern = r'(?m)^(?:\d+\.\s+|\•\s+|\-\s+|\*\s+)(.+)$'
                    bullet_points = re.findall(bullet_pattern, summary)
                    
                    if bullet_points:
                        for point in bullet_points:
                            bullet_para = doc.add_paragraph()
                            bullet_para.style = 'List Bullet'
                            bullet_para.add_run(point.strip()).font.size = Pt(11)
                    else:
                        # Fallback if no bullet points found
                        for line in summary.split('\n'):
                            if line.strip():
                                para = doc.add_paragraph()
                                para.add_run(line.strip()).font.size = Pt(11)
                    
                    # Add a separator
                    doc.add_paragraph()
                    
                    full_summary = f"Document Name: {document_name}\n\nKey Pointers:\n{summary}"
                    all_document_summaries.append(full_summary)
                    st.write(full_summary)
                    
                    file_progress.text(f"Completed {uploaded_file.name}")
                
                if partially_filtered_files:
                    st.warning(f"The following files had significant non-English content and were partially filtered: {', '.join(partially_filtered_files)}")
                
                # Save and provide download for Word document
                doc_output_path = "circulars_consolidated_summary.docx"
                doc.save(doc_output_path)
                
                with open(doc_output_path, "rb") as doc_file:
                    st.download_button(
                        "Download Summaries DOCX",
                        doc_file,
                        file_name="circulars_consolidated_summary.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
        except Exception as e:
            st.error(f"Error: {str(e)}")
            st.info("Please check that your API key is correct and that you have access to the selected model.")
elif summarize_button and (not uploaded_files or not api_key):
    if not uploaded_files:
        st.warning("Please upload PDF files before summarizing.")
    if not api_key:
        st.warning("Please enter your OpenAI API key before summarizing.")
else:
    st.info("Upload PDF files, enter your OpenAI API key, and click 'Summarize' to process documents.")
